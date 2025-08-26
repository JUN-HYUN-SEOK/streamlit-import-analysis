import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
import datetime
import os
import sys
import traceback
import openpyxl
from docx.enum.text import WD_COLOR_INDEX
import io
import zipfile
import time
from tempfile import NamedTemporaryFile

# 페이지 설정
st.set_page_config(
    page_title="수입신고 RISK 분석 시스템",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 메인 타이틀
col1, col2 = st.columns([4, 1])
with col1:
    st.title("📊 수입신고 RISK 분석 시스템")
with col2:
    st.markdown("<br><small style='color: #666; font-size: 0.8em;'>Made by 전자동</small>", unsafe_allow_html=True)
st.markdown("---")

# 사이드바 설정
st.sidebar.title("분석 옵션")
st.sidebar.markdown("분석할 엑셀 파일을 업로드하고 원하는 분석을 선택하세요.")

def read_excel_file(uploaded_file, progress_bar=None, status_text=None):
    """업로드된 엑셀 파일 읽기"""
    try:
        if status_text:
            status_text.text("📂 엑셀 파일 로드 중...")
        if progress_bar:
            progress_bar.progress(20)
        
        df = pd.read_excel(uploaded_file)
        
        if status_text:
            status_text.text(f"📊 데이터 로드 완료: {len(df):,}행, {len(df.columns)}열")
        if progress_bar:
            progress_bar.progress(40)
        
        df.columns = df.columns.str.strip()  # 컬럼 이름의 공백 제거
        
        if status_text:
            status_text.text("🔧 중복 컬럼명 처리 중...")
        if progress_bar:
            progress_bar.progress(50)
        
        # 중복된 컬럼명 처리
        cols = pd.Series(df.columns)
        duplicate_count = 0
        
        # 중복된 컬럼명이 있는지 확인
        duplicated_cols = cols[cols.duplicated()].unique()
        
        if len(duplicated_cols) > 0:
            # 중복된 각 컬럼에 대해 처리
            for dup in duplicated_cols:
                # 해당 컬럼이 나타나는 모든 인덱스 찾기
                dup_indices = cols[cols == dup].index.tolist()
                # 첫 번째는 그대로 두고, 나머지에 번호 추가
                for i, idx in enumerate(dup_indices):
                    if i > 0:  # 첫 번째가 아닌 경우에만 번호 추가
                        cols.iloc[idx] = f"{dup}_{i}"
                duplicate_count += 1
            
            # 변경된 컬럼명 적용
            df.columns = cols.tolist()
        
        if duplicate_count > 0 and status_text:
            status_text.text(f"⚠️ {duplicate_count}개의 중복 컬럼명 처리 완료")
        
        if progress_bar:
            progress_bar.progress(70)
        
        if status_text:
            status_text.text("🏷️ 컬럼 매핑 중...")
        
        # 1. 컬럼 존재 여부 확인 및 안전한 컬럼 매핑
        try:
            # 먼저 필요한 컬럼이 이미 있는지 확인
            has_rate_type = '세율구분' in df.columns
            has_tariff_rate = '관세실행세율' in df.columns
            
            if not has_rate_type or not has_tariff_rate:
                # 컬럼 인덱스 기반 매핑 시도
                column_list = df.columns.tolist()
                
                if len(column_list) > 71 and not has_tariff_rate:
                    # 71번째 컬럼을 관세실행세율로 매핑
                    if column_list[71] not in ['세율구분', '관세실행세율']:
                        df.rename(columns={column_list[71]: '관세실행세율'}, inplace=True)
                        has_tariff_rate = True
                
                if len(column_list) > 70 and not has_rate_type:
                    # 70번째 컬럼을 세율구분으로 매핑
                    if column_list[70] not in ['세율구분', '관세실행세율']:
                        df.rename(columns={column_list[70]: '세율구분'}, inplace=True)
                        has_rate_type = True
            
            # 없는 컬럼들은 기본값으로 생성
            if not has_rate_type:
                df['세율구분'] = 'A'
            if not has_tariff_rate:
                df['관세실행세율'] = 0
                
        except Exception as col_error:
            if status_text:
                status_text.text(f"⚠️ 컬럼 매핑 오류: 기본값으로 설정")
            # 기본 컬럼들 생성
            if '세율구분' not in df.columns:
                df['세율구분'] = 'A'
            if '관세실행세율' not in df.columns:
                df['관세실행세율'] = 0
        
        if progress_bar:
            progress_bar.progress(90)
        
        if status_text:
            status_text.text("🔢 데이터 타입 변환 중...")
        
        # 2. 관세실행세율 컬럼을 숫자형으로 변환
        try:
            if '관세실행세율' in df.columns:
                # 안전한 숫자형 변환
                tariff_col = df['관세실행세율']
                
                # 이미 숫자형인 경우 그대로 사용
                if pd.api.types.is_numeric_dtype(tariff_col):
                    df['관세실행세율'] = tariff_col.fillna(0)
                else:
                    # 문자열인 경우 숫자로 변환 시도
                    df['관세실행세율'] = pd.to_numeric(
                        tariff_col.astype(str).str.replace(',', '').fillna('0'), 
                        errors='coerce'
                    ).fillna(0)
        except Exception as convert_error:
            if status_text:
                status_text.text("⚠️ 숫자 변환 오류: 기본값 사용")
            df['관세실행세율'] = 0
        
        if progress_bar:
            progress_bar.progress(100)
        
        if status_text:
            status_text.text("✅ 데이터 처리 완료!")
        
        return df
    except Exception as e:
        if status_text:
            status_text.text(f"❌ 오류 발생: {str(e)}")
        st.error(f"엑셀 파일 읽기 실패: {str(e)}")
        st.error("파일 형식을 확인하거나 다른 파일을 시도해보세요.")
        return None

def process_data(df):
    """데이터 전처리"""
    try:
        # 컬럼 이름의 공백 제거
        df.columns = df.columns.str.strip()
        
        # 필요한 컬럼이 있는지 확인
        required_columns = ['관세실행세율', '세율구분']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            st.warning(f"누락된 컬럼: {missing_columns}")
            return None

        # 0% Risk 조건에 맞는 데이터 필터링
        df_zero_risk = df[
            (df['관세실행세율'] < 8) & 
            (~df['세율구분'].astype(str).str.match(r'^F.{3}$'))  # F로 시작하는 4자리 코드 제외
        ]

        # '세율구분'이 4자리인 행 제외
        df_filtered = df_zero_risk[df_zero_risk['세율구분'].apply(lambda x: len(str(x)) != 4)]

        return df_filtered
        
    except Exception as e:
        st.error(f"데이터 전처리 중 오류 발생: {e}")
        return None

def create_eight_percent_refund_analysis(df):
    """8% 환급 검토 분석"""
    try:
        # 필요한 컬럼만 선택
        selected_columns = [
            '수입신고번호',
            '수리일자',
            'B/L번호',
            '세번부호', 
            '세율구분',
            '세율설명',
            '관세실행세율',
            '적출국코드',
            '원산지코드',
            'FTA사후환급 검토',
            '규격1',
            '규격2',
            '규격3',
            '성분1',
            '성분2',
            '성분3',
            '실제관세액',
            '결제방법',
            '결제통화단위',
            '무역거래처상호',
            '무역거래처국가코드',
            '거래품명',
            '란번호',
            '행번호',
            '수량_1',
            '수량단위_1',
            '단가',
            '금액',
            '란결제금액',
            '행별관세'
        ]
        
        # 존재하는 컬럼만 선택
        base_columns = [col for col in selected_columns 
                       if col not in ['행별관세', 'FTA사후환급 검토'] and col in df.columns]
        
        # 원본 데이터를 복사하여 사용
        df_work = df[base_columns].copy()
        
        # 데이터 전처리
        df_work['세율구분'] = df_work['세율구분'].astype(str).str.strip()
        df_work['관세실행세율'] = pd.to_numeric(
            df_work['관세실행세율'].fillna(0), errors='coerce'
        ).fillna(0)
        df_work['실제관세액'] = pd.to_numeric(
            df_work['실제관세액'].fillna(0), errors='coerce'
        ).fillna(0)
        
        # 행별관세 계산에 필요한 컬럼들 전처리
        if '금액' in df_work.columns:
            df_work['금액'] = pd.to_numeric(
                df_work['금액'].fillna(0), errors='coerce'
            ).fillna(0)
        
        if '란결제금액' in df_work.columns:
            df_work['란결제금액'] = pd.to_numeric(
                df_work['란결제금액'].fillna(0), errors='coerce'
            ).fillna(0)
        
        # 행별관세 계산: (실제관세액 × 금액) ÷ 란결제금액
        if all(col in df_work.columns for col in ['실제관세액', '금액', '란결제금액']):
            df_work['행별관세'] = np.where(
                df_work['란결제금액'] != 0,
                (df_work['실제관세액'] * df_work['금액']) / df_work['란결제금액'],
                0
            )
        else:
            df_work['행별관세'] = 0
        
        # FTA사후환급 검토 컬럼 계산
        if '적출국코드' in df_work.columns and '원산지코드' in df_work.columns:
            df_work['FTA사후환급 검토'] = df_work.apply(
                lambda row: 'FTA사후환급 검토' if (
                    pd.notna(row['적출국코드']) and 
                    pd.notna(row['원산지코드']) and 
                    str(row['적출국코드']).strip() == str(row['원산지코드']).strip() and
                    str(row['적출국코드']).strip() != '' and
                    str(row['원산지코드']).strip() != ''
                ) else '', 
                axis=1
            )
        else:
            df_work['FTA사후환급 검토'] = ''
        
        # NaN 값을 0으로 대체
        df_work.fillna(0, inplace=True)
        df_work = df_work.infer_objects(copy=False)
        
        # 필터링 조건 적용
        df_filtered = df_work[
            (df_work['세율구분'] == 'A') & 
            (df_work['관세실행세율'] >= 8)
        ]
        
        # 최종 컬럼 순서 정리 (란결제금액은 계산 후 제거)
        final_columns = [col for col in selected_columns 
                        if col in df_filtered.columns and col != '란결제금액']
        df_filtered = df_filtered[final_columns]
        
        return df_filtered
        
    except Exception as e:
        st.error(f"8% 환급 검토 분석 중 오류 발생: {str(e)}")
        return None

def create_zero_percent_risk_analysis(df):
    """0% Risk 분석"""
    try:
        # 필요한 컬럼만 선택
        selected_columns = [
            '수입신고번호',
            '수리일자',
            'B/L번호',
            '세번부호', 
            '세율구분',
            '관세실행세율',
            '규격1',
            '규격2',
            '성분1',
            '실제관세액',
            '거래품명',
            '란번호',
            '행번호',
            '수량_1',
            '수량단위_1',
            '단가',
            '금액',
            '란결제금액',
            '행별관세'
        ]
        
        # 0% Risk 조건에 맞는 데이터 필터링
        df_zero_risk = df[
            (df['관세실행세율'] < 8) & 
            (~df['세율구분'].astype(str).str.match(r'^F.{3}$'))
        ]
        
        # 존재하는 컬럼만 선택
        base_columns = [col for col in selected_columns 
                       if col not in ['행별관세'] and col in df_zero_risk.columns]
        
        # 필요한 컬럼만 선택
        df_zero_risk = df_zero_risk[base_columns].copy()
        
        # 행별관세 계산에 필요한 컬럼들 전처리
        if '실제관세액' in df_zero_risk.columns:
            df_zero_risk['실제관세액'] = pd.to_numeric(
                df_zero_risk['실제관세액'].fillna(0), errors='coerce'
            ).fillna(0)
        
        if '금액' in df_zero_risk.columns:
            df_zero_risk['금액'] = pd.to_numeric(
                df_zero_risk['금액'].fillna(0), errors='coerce'
            ).fillna(0)
        
        if '란결제금액' in df_zero_risk.columns:
            df_zero_risk['란결제금액'] = pd.to_numeric(
                df_zero_risk['란결제금액'].fillna(0), errors='coerce'
            ).fillna(0)
        
        # 행별관세 계산: (실제관세액 × 금액) ÷ 란결제금액
        if all(col in df_zero_risk.columns for col in ['실제관세액', '금액', '란결제금액']):
            df_zero_risk['행별관세'] = np.where(
                df_zero_risk['란결제금액'] != 0,
                (df_zero_risk['실제관세액'] * df_zero_risk['금액']) / df_zero_risk['란결제금액'],
                0
            )
        else:
            df_zero_risk['행별관세'] = 0
        
        # NaN 값을 0으로 대체
        df_zero_risk.fillna(0, inplace=True)
        df_zero_risk = df_zero_risk.infer_objects(copy=False)
        
        # 최종 컬럼 순서 정리 (란결제금액은 계산 후 제거)
        final_columns = [col for col in selected_columns 
                        if col in df_zero_risk.columns and col != '란결제금액']
        df_zero_risk = df_zero_risk[final_columns]
        
        return df_zero_risk
    
    except Exception as e:
        st.error(f"0% Risk 분석 중 오류 발생: {str(e)}")
        return None

def create_tariff_risk_analysis(df):
    """세율 Risk 분석"""
    try:
        required_columns = [
            '수입신고번호', 
            '수리일자',
            '규격1', '규격2', '규격3',
            '성분1', '성분2', '성분3',
            '세번부호', 
            '세율구분', 
            '세율설명',
            '과세가격달러',
            '실제관세액',
            '결제방법',
            '금액',
            '란결제금액'
        ]
        
        # 규격1별 세번부호 분석
        if '규격1' in df.columns and '세번부호' in df.columns:
            # 규격1별로 세번부호의 고유값 개수를 계산
            risk_specs = df.groupby('규격1')['세번부호'].nunique()
            
            # 세번부호가 2개 이상인 규격1만 선택
            risk_specs = risk_specs[risk_specs > 1]
        else:
            risk_specs = pd.Series(dtype='object')
        
        if len(risk_specs) == 0:
            return pd.DataFrame()
            
        # 규격1 기준으로 정렬
        if '규격1' in df.columns:
            # 존재하는 컬럼만 선택
            available_columns = [col for col in required_columns if col in df.columns]
            risk_data = df[df['규격1'].isin(risk_specs.index)][available_columns].copy()
            
            # 행별관세 계산에 필요한 컬럼들 전처리
            if '실제관세액' in risk_data.columns:
                risk_data['실제관세액'] = pd.to_numeric(
                    risk_data['실제관세액'].fillna(0), errors='coerce'
                ).fillna(0)
            
            if '금액' in risk_data.columns:
                risk_data['금액'] = pd.to_numeric(
                    risk_data['금액'].fillna(0), errors='coerce'
                ).fillna(0)
            
            if '란결제금액' in risk_data.columns:
                risk_data['란결제금액'] = pd.to_numeric(
                    risk_data['란결제금액'].fillna(0), errors='coerce'
                ).fillna(0)
            
            # 행별관세 계산: (실제관세액 × 금액) ÷ 란결제금액
            if all(col in risk_data.columns for col in ['실제관세액', '금액', '란결제금액']):
                risk_data['행별관세'] = np.where(
                    risk_data['란결제금액'] != 0,
                    (risk_data['실제관세액'] * risk_data['금액']) / risk_data['란결제금액'],
                    0
                )
            else:
                risk_data['행별관세'] = 0
            
            # 규격1, 세번부호 기준 정렬
            risk_data = risk_data.sort_values(['규격1', '세번부호']).fillna('')
            
            # 최종 컬럼 순서 정리 (란결제금액은 계산 후 제거)
            final_columns = [col for col in available_columns if col != '란결제금액']
            if '행별관세' not in final_columns:
                final_columns.append('행별관세')
            risk_data = risk_data[final_columns]
        else:
            risk_data = pd.DataFrame(columns=required_columns)
        
        return risk_data
        
    except Exception as e:
        st.error(f"세율 Risk 분석 중 오류 발생: {e}")
        return pd.DataFrame()

def create_price_risk_analysis(df):
    """단가 Risk 분석"""
    try:
        # 필요한 컬럼 체크
        required_columns = ['규격1', '세번부호', '거래구분', '결제방법', '수리일자', '수입신고번호',
                          '단가', '결제통화단위', '거래품명', 
                          '란번호', '행번호', '수량_1', '수량단위_1', '금액']
        
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            available_columns = [col for col in required_columns if col in df.columns]
            if '단가' not in available_columns:
                return pd.DataFrame()
        else:
            available_columns = required_columns
        
        # 단가를 숫자형으로 변환
        df_work = df.copy()
        df_work['단가'] = pd.to_numeric(df_work['단가'].fillna(0), errors='coerce').fillna(0)
        
        # 단가가 0보다 큰 데이터만 분석
        df_work = df_work[df_work['단가'] > 0]
        
        if len(df_work) == 0:
            return pd.DataFrame()
        
        # 그룹화 기준 (규격1만 사용)
        group_columns = ['규격1']
        
        # 집계 함수 정의
        agg_dict = {
            '세번부호': 'first',
            '거래구분': 'first',
            '결제방법': 'first',
            '수리일자': ['min', 'max'],
            '수입신고번호': ['min', 'max'],
            '단가': ['mean', 'max', 'min', 'std', 'count'],
            '결제통화단위': 'first',
            '거래품명': 'first',
            '란번호': 'first',
            '행번호': 'first',
            '수량_1': 'first',
            '수량단위_1': 'first',
            '금액': 'sum'
        }
        
        # 존재하는 컬럼만 선택
        available_group_columns = [col for col in group_columns if col in df_work.columns]
        available_agg_dict = {col: agg_dict[col] for col in agg_dict if col in df_work.columns}
        
        grouped = df_work.groupby(available_group_columns).agg(available_agg_dict).reset_index()
        
        # 집계 후 컬럼명 재설정
        grouped_columns = list(grouped.columns)
        new_columns = []
        for col in grouped_columns:
            if isinstance(col, tuple):
                if col[0] == '단가' and col[1] == 'mean':
                    new_columns.append('평균단가')
                elif col[0] == '단가' and col[1] == 'max':
                    new_columns.append('최고단가')
                elif col[0] == '단가' and col[1] == 'min':
                    new_columns.append('최저단가')
                elif col[0] == '단가' and col[1] == 'std':
                    new_columns.append('단가표준편차')
                elif col[0] == '단가' and col[1] == 'count':
                    new_columns.append('데이터수')
                elif col[0] == '수리일자' and col[1] == 'min':
                    new_columns.append('Min 수리일자')
                elif col[0] == '수리일자' and col[1] == 'max':
                    new_columns.append('Max 수리일자')
                elif col[0] == '수입신고번호' and col[1] == 'min':
                    new_columns.append('Min 신고번호')
                elif col[0] == '수입신고번호' and col[1] == 'max':
                    new_columns.append('Max 신고번호')
                else:
                    if col[1] == 'first':
                        new_columns.append(col[0])
                    elif col[1] == 'sum':
                        new_columns.append(col[0])
                    else:
                        new_columns.append(f'{col[0]}_{col[1]}')
            else:
                new_columns.append(col)
        grouped.columns = new_columns
        
        # 위험도 계산
        grouped['단가편차율'] = np.where(
            grouped['평균단가'] > 0,
            (grouped['최고단가'] - grouped['최저단가']) / grouped['평균단가'],
            0
        )
        
        # 위험도 분류
        def classify_risk(row):
            if row['평균단가'] == 0:
                return '확인필요'
            elif row['단가편차율'] > 0.5:  # 50% 이상 편차
                return '매우높음'
            elif row['단가편차율'] > 0.3:  # 30% 이상 편차
                return '높음'
            elif row['단가편차율'] > 0.1:  # 10% 이상 편차
                return '보통'
            else:
                return '낮음'
        
        grouped['위험도'] = grouped.apply(classify_risk, axis=1)
        
        # 비고 생성
        grouped['비고'] = grouped.apply(lambda row: 
            f'평균단가 확인 필요' if row['평균단가'] == 0 
            else f'단가편차: {row["단가편차율"]*100:.1f}%', axis=1
        )
        
        return grouped
        
    except Exception as e:
        st.error(f"단가 Risk 분석 중 오류 발생: {str(e)}")
        return pd.DataFrame()

def create_summary_analysis(df_original):
    """Summary 분석"""
    try:
        summary_data = {}
        
        # 1. 전체 신고 건수
        if '수입신고번호' in df_original.columns:
            total_declarations = df_original['수입신고번호'].nunique()
        else:
            total_declarations = len(df_original)
        summary_data['전체 신고 건수'] = total_declarations
        
        # 2. 거래구분별 분석
        if '거래구분' in df_original.columns and '수입신고번호' in df_original.columns:
            trade_type_analysis = pd.pivot_table(df_original, 
                index=['거래구분'],
                values='수입신고번호',
                aggfunc='nunique',
                margins=True,
                margins_name='총계'
            ).reset_index()
        else:
            trade_type_analysis = pd.DataFrame({
                '거래구분': ['데이터 없음'],
                '수입신고번호': [0]
            })
        
        # 3. 세율구분별 분석
        if '세율구분' in df_original.columns and '수입신고번호' in df_original.columns:
            rate_type_analysis = pd.pivot_table(df_original,
                index='세율구분',
                values='수입신고번호',
                aggfunc='nunique'
            ).reset_index()
            # 총계 추가
            total_row = {'세율구분': '총계', '수입신고번호': rate_type_analysis['수입신고번호'].sum()}
            rate_type_analysis = pd.concat([rate_type_analysis, pd.DataFrame([total_row])], ignore_index=True)
        else:
            rate_type_analysis = pd.DataFrame({
                '세율구분': ['데이터 없음'],
                '수입신고번호': [0]
            })
        
        # 4. Risk 분석 요약
        if all(col in df_original.columns for col in ['관세실행세율', '세율구분', '수입신고번호']):
            zero_risk_df = df_original[
                (df_original['관세실행세율'] < 8) & 
                (~df_original['세율구분'].astype(str).str.match(r'^F.{3}$'))
            ]
            zero_risk_count = zero_risk_df['수입신고번호'].nunique()
            
            eight_percent_df = df_original[
                (df_original['세율구분'] == 'A') & 
                (df_original['관세실행세율'] >= 8)
            ]
            eight_percent_count = eight_percent_df['수입신고번호'].nunique()
        else:
            zero_risk_count = 0
            eight_percent_count = 0
        
        risk_analysis = pd.DataFrame({
            'Risk 유형': ['0% Risk', '8% 환급 검토'],
            '신고건수': [zero_risk_count, eight_percent_count],
            '비율(%)': [
                zero_risk_count/total_declarations*100 if total_declarations > 0 else 0,
                eight_percent_count/total_declarations*100 if total_declarations > 0 else 0
            ]
        })
        
        summary_data['거래구분별'] = trade_type_analysis
        summary_data['세율구분별'] = rate_type_analysis
        summary_data['Risk분석'] = risk_analysis
        
        return summary_data
        
    except Exception as e:
        st.error(f"Summary 분석 중 오류 발생: {str(e)}")
        return {}

def create_verification_methods_excel_sheet(writer):
    """검증방법 시트 생성 (엑셀용)"""
    try:
        # 워크시트 생성
        worksheet = writer.book.add_worksheet('검증방법')
        workbook = writer.book
        
        # 포맷 설정
        title_format = workbook.add_format({
            'font_name': 'Arial',
            'font_size': 14,
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'bg_color': '#4472C4',
            'font_color': 'white',
            'border': 1
        })
        
        subtitle_format = workbook.add_format({
            'font_name': 'Arial',
            'font_size': 12,
            'bold': True,
            'align': 'left',
            'valign': 'vcenter',
            'bg_color': '#D9E1F2',
            'border': 1
        })
        
        content_format = workbook.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'align': 'left',
            'valign': 'top',
            'border': 1,
            'text_wrap': True
        })
        
        highlight_format = workbook.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'align': 'left',
            'valign': 'top',
            'border': 1,
            'text_wrap': True,
            'bg_color': '#FFFF00'  # 노란색 배경
        })
        
        # 열 너비 설정
        worksheet.set_column(0, 0, 25)  # A열 - 시트명
        worksheet.set_column(1, 1, 60)  # B열 - 검증로직
        worksheet.set_column(2, 2, 40)  # C열 - 특이사항
        
        current_row = 0
        
        # 제목
        worksheet.merge_range(current_row, 0, current_row, 2, '수입신고 분석 검증방법', title_format)
        worksheet.set_row(current_row, 30)
        current_row += 2
        
        # 1. 8% 환급 검토
        worksheet.write(current_row, 0, '1. 8% 환급 검토', subtitle_format)
        worksheet.write(current_row, 1, 
            '• 필터링 조건: 세율구분 = "A" AND 관세실행세율 ≥ 8%\n' +
            '• 목적: 8% 환급 검토가 필요한 수입신고 건들 식별\n' +
            '• 추가 컬럼: 적출국코드, 원산지코드, 무역거래처상호, 무역거래처국가코드\n' +
            '• 행별관세 계산: (실제관세액 × 금액) ÷ 란결제금액', 
            content_format)
        worksheet.write(current_row, 2, 
            '• 세율구분 "A"는 일반적으로 가장 관세율이 높은 구분\n' +
            '• 8% 이상의 관세율은 환급 대상이 될 수 있음\n' +
            '• FTA사후환급 검토: 적출국=원산지인 경우 표시', 
            highlight_format)
        worksheet.set_row(current_row, 80)
        current_row += 1
        
        # 2. 0% Risk
        worksheet.write(current_row, 0, '2. 0% Risk', subtitle_format)
        worksheet.write(current_row, 1, 
            '• 필터링 조건: 관세실행세율 < 8% AND 세율구분 ≠ F***\n' +
            '• 목적: 관세율이 낮거나 면세 대상이지만 추가 검토가 필요한 건들\n' +
            '• F로 시작하는 4자리 코드는 특별한 세율구분으로 제외\n' +
            '• 행별관세 계산: (실제관세액 × 금액) ÷ 란결제금액', 
            content_format)
        worksheet.write(current_row, 2, 
            '• 관세율이 낮은데도 특별한 세율구분이 아닌 경우 주의 필요\n' +
            '• 면세 대상이지만 실제로는 관세가 부과될 수 있는 경우\n' +
            '• 관세실행세율이 0%인 경우 노란색으로 강조 표시', 
            highlight_format)
        worksheet.set_row(current_row, 80)
        current_row += 1
        
        # 3. 세율 Risk
        worksheet.write(current_row, 0, '3. 세율 Risk', subtitle_format)
        worksheet.write(current_row, 1, 
            '• 분석 방법: 규격1 기준으로 그룹화하여 세번부호의 고유값 개수 확인\n' +
            '• 위험 판정: 동일 규격1에 대해 서로 다른 세번부호가 2개 이상인 경우\n' +
            '• 목적: 동일 상품(규격1)에 대한 세번부호 불일치 위험 식별\n' +
            '• 예시: "DEMO SYS 1ML LG 0000-S000P1MLF"에 여러 세번부호 적용\n' +
            '• 행별관세 계산: (실제관세액 × 금액) ÷ 란결제금액', 
            content_format)
        worksheet.write(current_row, 2, 
            '• 동일 상품인데 다른 세번부호가 적용되면 관세율 차이 발생\n' +
            '• 세번부호 분류 오류 가능성 또는 상품 특성 차이\n' +
            '• 세율 Risk 발견 시 해당 규격1의 세번부호들을 상세 검토 필요\n' +
            '• 세번부호가 다른 경우 노란색으로 강조 표시', 
            highlight_format)
        worksheet.set_row(current_row, 100)
        current_row += 1
        
        # 4. 단가 Risk
        worksheet.write(current_row, 0, '4. 단가 Risk', subtitle_format)
        worksheet.write(current_row, 1, 
            '• 그룹화 기준: 규격1\n' +
            '• 위험도 계산: 단가편차율 = (최고단가 - 최저단가) ÷ 평균단가\n' +
            '• 위험도 분류:\n' +
            '  - 10% 초과~30% 이하: "보통"\n' +
            '  - 30% 초과~50% 이하: "높음"\n' +
            '  - 50% 초과: "매우높음"\n' +
            '• 특이사항: 평균단가가 0인 경우 "확인필요"로 분류\n' +
            '• 추가 정보: Min/Max 신고번호의 수리일자 표시', 
            content_format)
        worksheet.write(current_row, 2, 
            '• 단가 변동성이 10% 초과하면 주의 필요\n' +
            '• 30% 초과는 높은 위험, 50% 초과는 매우 비정상적인 가격 차이\n' +
            '• 평균단가 0은 데이터 오류 또는 특별한 거래 형태\n' +
            '• 수리일자 차이로 시간적 변동성 확인 가능\n' +
            '• 위험도가 "높음", "매우높음", "확인필요"인 경우 노란색 강조', 
            highlight_format)
        worksheet.set_row(current_row, 120)
        current_row += 1
        
        # 5. Summary
        worksheet.write(current_row, 0, '5. Summary', subtitle_format)
        worksheet.write(current_row, 1, 
            '• 전체 신고 건수: 수입신고번호 기준 고유 건수\n' +
            '• 거래구분별 분석: 거래구분별 신고건수 피벗 테이블\n' +
            '• 세율구분별 분석: 세율구분별 신고건수 및 비중\n' +
            '• Risk 분석 요약: 0% Risk와 8% 환급 검토 건수 및 비율\n' +
            '• 세번부호별 세율구분 및 실행세율 분석', 
            content_format)
        worksheet.write(current_row, 2, 
            '• 전체적인 수입신고 현황 파악\n' +
            '• Risk 분포를 통한 우선순위 설정 가능\n' +
            '• 차트와 그래프로 시각적 분석 제공', 
            highlight_format)
        worksheet.set_row(current_row, 80)
        current_row += 1
        
        # 6. 원본데이터
        worksheet.write(current_row, 0, '6. 원본데이터', subtitle_format)
        worksheet.write(current_row, 1, 
            '• 분석에 사용된 원본 엑셀 파일의 모든 데이터\n' +
            '• 상위 1000개 행만 표시 (파일 크기 제한)\n' +
            '• 모든 컬럼과 원본 데이터 구조 확인 가능\n' +
            '• 필터링 및 정렬 기능 제공\n' +
            '• 중복 컬럼명 자동 처리됨', 
            content_format)
        worksheet.write(current_row, 2, 
            '• 원본 데이터와 분석 결과 비교 검토 가능\n' +
            '• 데이터 품질 및 구조 확인용\n' +
            '• 중복 컬럼은 _1, _2 등으로 구분', 
            highlight_format)
        worksheet.set_row(current_row, 80)
        current_row += 1
        
        # 특이사항 표시 방법
        worksheet.write(current_row, 0, '특이사항 표시 방법', subtitle_format)
        worksheet.write(current_row, 1, 
            '• 노란색 배경: 각 시트에서 특별히 주의가 필요한 항목\n' +
            '• 8% 환급 검토: 관세실행세율 8% 이상, FTA사후환급 검토 대상\n' +
            '• 0% Risk: 관세실행세율이 0%인 경우\n' +
            '• 세율 Risk: 동일 규격1에 다른 세번부호 적용\n' +
            '• 단가 Risk: 위험도 "높음", "매우높음", "확인필요"\n' +
            '• Summary: 세율구분/실행세율 종류수가 2개 이상인 세번부호', 
            content_format)
        worksheet.write(current_row, 2, 
            '• 노란색으로 표시된 항목은 반드시 검토 필요\n' +
            '• 데이터 오류 또는 비정상적인 거래 형태일 가능성\n' +
            '• 세관 신고 시 추가 확인이 필요한 항목들\n' +
            '• Made by 전자동 (Wooshin Customs Broker)', 
            highlight_format)
        worksheet.set_row(current_row, 100)
        
        # 페이지 설정
        worksheet.set_header('&C&B검증방법')
        worksheet.set_footer('&R&D &T')
        
        return True
        
    except Exception as e:
        print(f"검증방법 시트 생성 중 오류 발생: {str(e)}")
        return False

def create_excel_file(df_original, eight_percent_data, zero_risk_data, tariff_risk_data, price_risk_data, summary_data):
    """엑셀 파일 생성"""
    try:
        # 메모리에서 엑셀 파일 생성
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            workbook = writer.book
            
            # 포맷 설정
            header_format = workbook.add_format({
                'bold': True,
                'bg_color': '#D9E1F2',
                'border': 1,
                'align': 'center'
            })
            
            # Summary 시트
            if summary_data:
                summary_sheet = workbook.add_worksheet('Summary')
                row = 0
                
                # 제목
                summary_sheet.merge_range(row, 0, row, 3, '수입신고 분석 보고서', 
                                        workbook.add_format({'bold': True, 'font_size': 16, 'align': 'center'}))
                row += 2
                
                # 전체 신고 건수
                summary_sheet.write(row, 0, '전체 신고 건수', header_format)
                summary_sheet.write(row, 1, summary_data.get('전체 신고 건수', 0))
                row += 2
                
                # 거래구분별
                if '거래구분별' in summary_data:
                    summary_sheet.write(row, 0, '거래구분별 분석', header_format)
                    row += 1
                    summary_data['거래구분별'].to_excel(writer, sheet_name='Summary', startrow=row, startcol=0, index=False)
                    row += len(summary_data['거래구분별']) + 2
                
                # 세율구분별
                if '세율구분별' in summary_data:
                    summary_sheet.write(row, 0, '세율구분별 분석', header_format)
                    row += 1
                    summary_data['세율구분별'].to_excel(writer, sheet_name='Summary', startrow=row, startcol=0, index=False)
                    row += len(summary_data['세율구분별']) + 2
                
                # Risk 분석
                if 'Risk분석' in summary_data:
                    summary_sheet.write(row, 0, 'Risk 분석 요약', header_format)
                    row += 1
                    summary_data['Risk분석'].to_excel(writer, sheet_name='Summary', startrow=row, startcol=0, index=False)
            
            # 8% 환급 검토 시트
            if not eight_percent_data.empty:
                eight_percent_data.to_excel(writer, sheet_name='8% 환급 검토', index=False)
            
            # 0% Risk 시트
            if not zero_risk_data.empty:
                zero_risk_data.to_excel(writer, sheet_name='0% Risk', index=False)
            
            # 세율 Risk 시트
            if not tariff_risk_data.empty:
                tariff_risk_data.to_excel(writer, sheet_name='세율 Risk', index=False)
            
            # 단가 Risk 시트
            if not price_risk_data.empty:
                price_risk_data.to_excel(writer, sheet_name='단가 Risk', index=False)
            
            # 원본데이터 시트 (상위 1000개 행만)
            max_rows = min(1000, len(df_original))
            df_original.head(max_rows).to_excel(writer, sheet_name='원본데이터', index=False)
            
            # 검증방법 시트 생성
            create_verification_methods_excel_sheet(writer)
        
        output.seek(0)
        return output.getvalue()
        
    except Exception as e:
        st.error(f"엑셀 파일 생성 중 오류 발생: {str(e)}")
        return None

def create_word_document(eight_percent_data, zero_risk_data, tariff_risk_data, price_risk_data, summary_data):
    """워드 문서 생성"""
    try:
        doc = Document()
        
        # 제목 추가
        doc.add_heading('수입신고 분석 보고서', 0)
        
        # 날짜 추가
        doc.add_paragraph(datetime.datetime.now().strftime("%Y년 %m월 %d일"))
        
        # Summary 정보
        if summary_data:
            doc.add_heading('분석 요약', level=1)
            p = doc.add_paragraph()
            p.add_run(f"전체 신고 건수: {summary_data.get('전체 신고 건수', 0)}건").bold = True
            
            if 'Risk분석' in summary_data:
                risk_df = summary_data['Risk분석']
                p.add_run("\n\nRisk 분석 결과:")
                for _, row in risk_df.iterrows():
                    p.add_run(f"\n- {row['Risk 유형']}: {row['신고건수']}건 ({row['비율(%)']:.1f}%)")
        
        # 8% 환급 검토
        if not eight_percent_data.empty:
            doc.add_heading('8% 환급 검토', level=1)
            doc.add_paragraph(f'총 {len(eight_percent_data)}건의 8% 환급 검토 대상이 발견되었습니다.')
        
        # 0% Risk
        if not zero_risk_data.empty:
            doc.add_heading('0% Risk 분석', level=1)
            doc.add_paragraph(f'총 {len(zero_risk_data)}건의 0% Risk가 발견되었습니다.')
        
        # 세율 Risk
        if not tariff_risk_data.empty:
            doc.add_heading('세율 Risk 분석', level=1)
            doc.add_paragraph(f'총 {len(tariff_risk_data)}건의 세율 Risk가 발견되었습니다.')
        
        # 단가 Risk
        if not price_risk_data.empty:
            doc.add_heading('단가 Risk 분석', level=1)
            doc.add_paragraph(f'총 {len(price_risk_data)}건의 단가 Risk가 발견되었습니다.')
            
            # 위험도별 분포
            if '위험도' in price_risk_data.columns:
                risk_summary = price_risk_data['위험도'].value_counts()
                p = doc.add_paragraph("위험도 분포:")
                for risk, count in risk_summary.items():
                    p.add_run(f"\n- {risk}: {count}건")
        
        # 워드 파일을 메모리에서 생성
        doc_output = io.BytesIO()
        doc.save(doc_output)
        doc_output.seek(0)
        return doc_output.getvalue()
        
    except Exception as e:
        st.error(f"워드 문서 생성 중 오류 발생: {str(e)}")
        return None

# 메인 애플리케이션
def main():
    # 파일 업로드
    uploaded_file = st.file_uploader(
        "📁 엑셀 파일 업로드", 
        type=['xlsx', 'xls'],
        help="분석할 수입신고 데이터가 포함된 엑셀 파일을 업로드하세요."
    )
    
    if uploaded_file is not None:
        try:
            # 파일 정보 표시
            st.success(f"✅ 파일 업로드 완료: {uploaded_file.name}")
            
            # 데이터 읽기
            progress_container = st.container()
            with progress_container:
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                status_text.text("📊 엑셀 파일 읽기 시작...")
                progress_bar.progress(10)
                
                df_original = read_excel_file(uploaded_file, progress_bar, status_text)
                
                # 잠시 완료 메시지 표시 후 정리
                time.sleep(1)
                progress_bar.empty()
                status_text.empty()
                
            if df_original is not None:
                st.success(f"📈 데이터 로드 완료: {len(df_original):,}건의 데이터")
                
                # 데이터 미리보기
                with st.expander("📋 데이터 미리보기"):
                    try:
                        # 데이터프레임 표시 전에 문제가 될 수 있는 컬럼 타입 정리
                        df_preview = df_original.head(10).copy()
                        
                        # 모든 컬럼을 문자열로 변환하여 표시 (미리보기용)
                        for col in df_preview.columns:
                            df_preview[col] = df_preview[col].astype(str)
                        
                        st.dataframe(df_preview, use_container_width=True)
                        st.info(f"총 {len(df_original):,}행, {len(df_original.columns)}열")
                        
                        # 중복 컬럼이 있었는지 표시
                        duplicate_cols = [col for col in df_original.columns if '_1' in col or '_2' in col]
                        if duplicate_cols:
                            st.warning(f"중복된 컬럼명이 감지되어 자동으로 처리되었습니다: {', '.join(duplicate_cols[:5])}")
                            
                    except Exception as preview_error:
                        st.error(f"데이터 미리보기 중 오류: {preview_error}")
                        st.info(f"데이터는 정상적으로 로드되었습니다. 총 {len(df_original):,}행, {len(df_original.columns)}열")
                
                # 분석 옵션 선택
                st.sidebar.markdown("### 분석 옵션")
                analysis_options = st.sidebar.multiselect(
                    "수행할 분석을 선택하세요:",
                    ["Summary", "8% 환급 검토", "0% Risk", "세율 Risk", "단가 Risk"],
                    default=["Summary", "8% 환급 검토", "0% Risk", "세율 Risk", "단가 Risk"]
                )
                
                if st.sidebar.button("🔍 분석 시작", type="primary"):
                    results = {}
                    
                    # 각 분석 수행
                    analysis_container = st.container()
                    with analysis_container:
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        total_analyses = len(analysis_options)
                        current_step = 0
                        
                        status_text.text("🚀 분석을 시작합니다...")
                        progress_bar.progress(0)
                    
                        # Summary 분석
                        if "Summary" in analysis_options:
                            current_step += 1
                            status_text.text(f"📊 Summary 분석 중... ({current_step}/{total_analyses})")
                            progress_bar.progress(current_step / total_analyses)
                            results['summary'] = create_summary_analysis(df_original)
                        
                        # 8% 환급 검토
                        if "8% 환급 검토" in analysis_options:
                            current_step += 1
                            status_text.text(f"💰 8% 환급 검토 분석 중... ({current_step}/{total_analyses})")
                            progress_bar.progress(current_step / total_analyses)
                            results['eight_percent'] = create_eight_percent_refund_analysis(df_original)
                        
                        # 0% Risk
                        if "0% Risk" in analysis_options:
                            current_step += 1
                            status_text.text(f"🟢 0% Risk 분석 중... ({current_step}/{total_analyses})")
                            progress_bar.progress(current_step / total_analyses)
                            results['zero_risk'] = create_zero_percent_risk_analysis(df_original)
                        
                        # 세율 Risk
                        if "세율 Risk" in analysis_options:
                            current_step += 1
                            status_text.text(f"⚠️ 세율 Risk 분석 중... ({current_step}/{total_analyses})")
                            progress_bar.progress(current_step / total_analyses)
                            results['tariff_risk'] = create_tariff_risk_analysis(df_original)
                        
                        # 단가 Risk
                        if "단가 Risk" in analysis_options:
                            current_step += 1
                            status_text.text(f"💲 단가 Risk 분석 중... ({current_step}/{total_analyses})")
                            progress_bar.progress(current_step / total_analyses)
                            results['price_risk'] = create_price_risk_analysis(df_original)
                        
                        progress_bar.progress(1.0)
                        status_text.text("🎉 모든 분석이 완료되었습니다!")
                    
                    # 결과 표시
                    st.success("🎉 분석이 완료되었습니다!")
                    
                    # 탭으로 결과 표시
                    tab_names = []
                    tab_data = []
                    
                    if 'summary' in results and results['summary']:
                        tab_names.append("📊 Summary")
                        tab_data.append(('summary', results['summary']))
                    
                    if 'eight_percent' in results and not results['eight_percent'].empty:
                        tab_names.append("💰 8% 환급 검토")
                        tab_data.append(('eight_percent', results['eight_percent']))
                    
                    if 'zero_risk' in results and not results['zero_risk'].empty:
                        tab_names.append("🟢 0% Risk")
                        tab_data.append(('zero_risk', results['zero_risk']))
                    
                    if 'tariff_risk' in results and not results['tariff_risk'].empty:
                        tab_names.append("⚠️ 세율 Risk")
                        tab_data.append(('tariff_risk', results['tariff_risk']))
                    
                    if 'price_risk' in results and not results['price_risk'].empty:
                        tab_names.append("💲 단가 Risk")
                        tab_data.append(('price_risk', results['price_risk']))
                    
                    if tab_names:
                        tabs = st.tabs(tab_names)
                        
                        for i, (tab_type, data) in enumerate(tab_data):
                            with tabs[i]:
                                if tab_type == 'summary':
                                    st.subheader("분석 요약")
                                    
                                    col1, col2, col3 = st.columns(3)
                                    with col1:
                                        st.metric("전체 신고 건수", f"{data.get('전체 신고 건수', 0):,}건")
                                    
                                    if 'Risk분석' in data:
                                        risk_df = data['Risk분석']
                                        with col2:
                                            zero_risk = risk_df[risk_df['Risk 유형'] == '0% Risk']['신고건수'].iloc[0] if len(risk_df) > 0 else 0
                                            st.metric("0% Risk", f"{zero_risk:,}건")
                                        with col3:
                                            eight_percent = risk_df[risk_df['Risk 유형'] == '8% 환급 검토']['신고건수'].iloc[0] if len(risk_df) > 1 else 0
                                            st.metric("8% 환급 검토", f"{eight_percent:,}건")
                                    
                                    # 상세 분석 결과 표시
                                    if 'Risk분석' in data:
                                        st.subheader("Risk 분석 상세")
                                        try:
                                            st.dataframe(data['Risk분석'], use_container_width=True)
                                        except Exception as e:
                                            st.error(f"Risk 분석 표시 중 오류: {e}")
                                    
                                    if '거래구분별' in data:
                                        st.subheader("거래구분별 분석")
                                        try:
                                            st.dataframe(data['거래구분별'], use_container_width=True)
                                        except Exception as e:
                                            st.error(f"거래구분별 분석 표시 중 오류: {e}")
                                    
                                    if '세율구분별' in data:
                                        st.subheader("세율구분별 분석")
                                        try:
                                            st.dataframe(data['세율구분별'], use_container_width=True)
                                        except Exception as e:
                                            st.error(f"세율구분별 분석 표시 중 오류: {e}")
                                
                                else:
                                    # 데이터프레임 표시
                                    st.subheader(f"총 {len(data):,}건의 데이터")
                                    
                                    # 검색 기능
                                    search_term = st.text_input(f"{tab_names[i]} 검색", key=f"search_{tab_type}")
                                    
                                    try:
                                        if search_term:
                                            # 안전한 검색을 위해 모든 컬럼을 문자열로 변환
                                            data_str = data.astype(str)
                                            mask = data_str.apply(lambda x: x.str.contains(search_term, case=False, na=False)).any(axis=1)
                                            filtered_data = data[mask]
                                            st.write(f"검색 결과: {len(filtered_data)}건")
                                            
                                            # 검색 결과 표시
                                            if len(filtered_data) > 0:
                                                # 안전한 표시를 위해 문자열 변환
                                                display_data = filtered_data.copy()
                                                for col in display_data.columns:
                                                    display_data[col] = display_data[col].astype(str)
                                                st.dataframe(display_data, use_container_width=True)
                                            else:
                                                st.info("검색 결과가 없습니다.")
                                        else:
                                            # 페이지네이션
                                            page_size = 100
                                            total_pages = (len(data) - 1) // page_size + 1
                                            page = st.selectbox(f"페이지 ({total_pages}페이지 중)", range(1, total_pages + 1), key=f"page_{tab_type}")
                                            
                                            start_idx = (page - 1) * page_size
                                            end_idx = start_idx + page_size
                                            display_data = data.iloc[start_idx:end_idx].copy()
                                            
                                            # 안전한 표시를 위해 문자열 변환
                                            for col in display_data.columns:
                                                display_data[col] = display_data[col].astype(str)
                                            st.dataframe(display_data, use_container_width=True)
                                            
                                    except Exception as display_error:
                                        st.error(f"데이터 표시 중 오류: {display_error}")
                                        st.info("데이터 형식에 문제가 있어 표시할 수 없습니다. 분석은 정상적으로 완료되었습니다.")
                    
                    # 파일 다운로드
                    st.markdown("---")
                    st.subheader("📥 결과 파일 다운로드")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # 엑셀 파일 생성 및 다운로드
                        excel_data = create_excel_file(
                            df_original,
                            results.get('eight_percent', pd.DataFrame()),
                            results.get('zero_risk', pd.DataFrame()),
                            results.get('tariff_risk', pd.DataFrame()),
                            results.get('price_risk', pd.DataFrame()),
                            results.get('summary', {})
                        )
                        
                        if excel_data:
                            st.download_button(
                                label="📊 Excel 파일 다운로드",
                                data=excel_data,
                                file_name=f"수입신고분석_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                    
                    with col2:
                        # 워드 문서 생성 및 다운로드
                        word_data = create_word_document(
                            results.get('eight_percent', pd.DataFrame()),
                            results.get('zero_risk', pd.DataFrame()),
                            results.get('tariff_risk', pd.DataFrame()),
                            results.get('price_risk', pd.DataFrame()),
                            results.get('summary', {})
                        )
                        
                        if word_data:
                            st.download_button(
                                label="📄 Word 파일 다운로드",
                                data=word_data,
                                file_name=f"수입신고분석_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
            
        except Exception as e:
            st.error(f"❌ 오류가 발생했습니다: {str(e)}")
            
            # 사용자에게 친숙한 오류 메시지 제공
            error_message = str(e).lower()
            if "arg must be a list" in error_message:
                st.warning("💡 **해결 방법:** 엑셀 파일의 데이터 형식에 문제가 있을 수 있습니다.")
                st.info("다음을 확인해주세요:\n- 파일이 손상되지 않았는지\n- 빈 셀이나 특수문자가 많지 않은지\n- 다른 엑셀 파일로 테스트해보세요")
            elif "duplicate" in error_message:
                st.warning("💡 **해결 방법:** 중복된 컬럼명이 있습니다.")
                st.info("엑셀 파일의 헤더(첫 번째 행)에 같은 이름의 컬럼이 여러 개 있는지 확인해주세요.")
            elif "memory" in error_message or "size" in error_message:
                st.warning("💡 **해결 방법:** 파일이 너무 큽니다.")
                st.info("더 작은 데이터 파일로 테스트하거나, 데이터를 분할해서 업로드해보세요.")
            else:
                st.warning("💡 **일반적인 해결 방법:**")
                st.info("1. 파일이 .xlsx 또는 .xls 형식인지 확인\n2. 파일이 손상되지 않았는지 확인\n3. 다른 파일로 테스트\n4. 브라우저 새로고침 후 재시도")
            
            # 개발자를 위한 상세 정보 (접을 수 있는 형태)
            with st.expander("🔧 개발자 정보 (상세 오류)"):
                st.code(traceback.format_exc())
    
    else:
        # 사용법 안내
        st.info("👆 좌측 사이드바에서 엑셀 파일을 업로드해주세요.")
        
        with st.expander("ℹ️ 사용법 안내"):
            st.markdown("""
            ### 📋 사용 방법
            1. **파일 업로드**: 분석할 수입신고 데이터가 포함된 엑셀 파일을 업로드하세요.
            2. **분석 옵션 선택**: 사이드바에서 원하는 분석 유형을 선택하세요.
            3. **분석 실행**: '분석 시작' 버튼을 클릭하여 분석을 시작하세요.
            4. **결과 확인**: 탭에서 분석 결과를 확인하세요.
            5. **파일 다운로드**: Excel 및 Word 형태로 결과를 다운로드하세요.
            
            ### 📊 분석 유형
            - **Summary**: 전체적인 분석 요약 및 통계
            - **8% 환급 검토**: 8% 이상 관세율에 대한 환급 검토 대상
            - **0% Risk**: 낮은 관세율 Risk 분석
            - **세율 Risk**: 세번부호 불일치 위험 분석
            - **단가 Risk**: 단가 변동성 위험 분석
            
            ### 📁 지원 파일 형식
            - Excel 파일 (.xlsx, .xls)
            """)

if __name__ == "__main__":
    main()
    
    # 화면 하단에 회사명 표시
    st.markdown("---")
    st.markdown(
        "<div style='text-align: center; color: #888; font-size: 0.9em; padding: 20px;'>"
        "© Wooshin Customs Broker"
        "</div>", 
        unsafe_allow_html=True
    )
